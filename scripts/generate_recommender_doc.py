"""
generate_recommender_doc.py
────────────────────────────
Generates the Word documentation for the UBP Attrition Recommender System.
Run from project root:
    python scripts/generate_recommender_doc.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT     = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "UBP_Recommender_System_Documentation.docx"


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "1F4E79")
        shade.set(qn("w:color"), "FFFFFF")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shade)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


# ── Document ───────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("UBP Client Attrition API", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph("Rule-Based Product Recommender System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size  = Pt(14)
    sub.runs[0].font.bold  = True
    sub.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    sub2 = doc.add_paragraph("Technical Design & Business Logic Documentation")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].font.size   = Pt(11)
    sub2.runs[0].font.italic = True
    sub2.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph("UnionBank of the Philippines  ·  Data Science Team").runs[0].font.size = Pt(10)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 1. Overview ────────────────────────────────────────────────────────────
    set_heading(doc, "1. Overview", 1, (31, 78, 121))
    add_paragraph(doc, (
        "The recommender system is a rule-based engine embedded in the UBP Attrition "
        "API. It runs immediately after the attrition prediction pipeline and produces "
        "a prioritised list of up to five UnionBank product recommendations for each "
        "client. Its purpose is to give relationship managers a concrete, actionable "
        "next step based on the client's financial profile and predicted attrition risk."
    ))

    set_heading(doc, "1.1  Why Rule-Based and Not ML-Based?", 2)
    add_paragraph(doc, (
        "The EDA phase of the attrition model project established that all available "
        "features carry near-zero discriminative signal for churn prediction "
        "(Cohen's d < 0.03 for all continuous features; Cramér's V < 0.01 for all "
        "categorical features). The best model achieved a PR-AUC of 0.0495 — "
        "approximately equal to the random baseline of 0.05 on a 5% positive-rate dataset."
    ))
    add_paragraph(doc, (
        "Building a second ML model on top of an already weak signal would amplify "
        "noise rather than add value. Rule-based logic anchored on observable, "
        "auditable client attributes — income, card tier, tenure, engagement — "
        "is more interpretable, more defensible to bank governance, and more "
        "immediately actionable for a relationship manager."
    ))

    doc.add_paragraph()

    # ── 2. Architecture ────────────────────────────────────────────────────────
    set_heading(doc, "2. System Architecture", 1, (31, 78, 121))
    add_paragraph(doc, (
        "The recommender is entirely downstream of the attrition pipeline. "
        "It does not affect model training, feature engineering, or serialization. "
        "The data flow is:"
    ))
    add_code_block(doc,
        "Raw client features (11 fields)\n"
        "         │\n"
        "         ▼\n"
        "  pipeline.predict_proba()        ← sklearn inference pipeline\n"
        "         │\n"
        "         ▼\n"
        "  attrition_probability (float)\n"
        "         │\n"
        "         ├── + original raw features\n"
        "         ▼\n"
        "  get_recommendations()           ← api/recommender.py\n"
        "         │\n"
        "         ▼\n"
        "  RecommendationBundle            ← returned in API response"
    )
    add_paragraph(doc, (
        "The recommender receives the attrition probability and all 11 raw input "
        "features. It derives additional signals internally (income tier, risk tier, "
        "life stage, credit utilization) and applies five independent rule layers."
    ))

    doc.add_paragraph()

    # ── 3. Input Features ──────────────────────────────────────────────────────
    set_heading(doc, "3. Input Features Used by the Recommender", 1, (31, 78, 121))
    add_paragraph(doc, (
        "The following features are passed from the API to the recommender. "
        "Features marked 'Unused' are accepted for forward compatibility but "
        "intentionally excluded from rules — the EDA confirmed Cramér's V < 0.01 "
        "for Gender, MaritalStatus, EducationLevel, and Country, meaning they have "
        "no meaningful relationship with any outcome."
    ))
    doc.add_paragraph()
    add_table(doc,
        ["Feature", "Type", "Used In Rules", "Reason If Unused"],
        [
            ["attrition_probability", "float [0,1]", "Yes", ""],
            ["Age",                   "int",         "Yes", ""],
            ["Income",                "float (PHP/yr)", "Yes", ""],
            ["CreditLimit",           "float",       "Yes", ""],
            ["TotalTransactions",     "int",         "Yes", ""],
            ["TotalSpend",            "float",       "Yes", ""],
            ["Tenure",                "int (months)","Yes", ""],
            ["CardType",              "str",         "Yes", ""],
            ["Gender",                "str",         "No",  "Cramér's V < 0.01 in EDA"],
            ["EducationLevel",        "str",         "No",  "Cramér's V < 0.01 in EDA"],
            ["MaritalStatus",         "str",         "No",  "Cramér's V < 0.01 in EDA"],
            ["Country",               "str",         "No",  "Target-encoded SD = 0.0068"],
        ],
        col_widths=[1.6, 1.2, 1.1, 2.3]
    )

    doc.add_paragraph()

    # ── 4. Derived Signals ─────────────────────────────────────────────────────
    set_heading(doc, "4. Derived Signals", 1, (31, 78, 121))
    add_paragraph(doc, (
        "Before the rule layers execute, four signals are derived from the raw inputs. "
        "These are computed at runtime — not features from the ML pipeline."
    ))
    doc.add_paragraph()
    add_table(doc,
        ["Derived Signal", "Formula / Logic", "Purpose"],
        [
            ["income_tier",
             "Entry (<360k) / Standard (360k–1.2M) / Premium (1.2M–3M) / Wealth (>3M)",
             "Primary driver for product eligibility and cross-sell targeting"],
            ["risk_tier",
             "Critical (≥0.65) / High (≥0.45) / Moderate (≥0.25) / Low (<0.25)",
             "Overlay for retention rule activation. Acknowledged as weak signal."],
            ["utilization",
             "TotalSpend / CreditLimit",
             "Engagement health indicator. < 15% = underusing; > 85% = credit stress."],
            ["life_stage",
             "Age < 30: Young Professional; 30–39: Early Career; 40–54: Mid Career; 55+: Pre-Retirement",
             "Drives life-stage product recommendations (EON, home loan, BancAssurance)"],
        ],
        col_widths=[1.4, 2.8, 2.0]
    )

    doc.add_paragraph()

    # ── 5. Income Thresholds ───────────────────────────────────────────────────
    set_heading(doc, "5. Income Tier Thresholds", 1, (31, 78, 121))
    add_paragraph(doc, (
        "Thresholds are derived from Bangko Sentral ng Pilipinas (BSP) consumer "
        "banking income segmentation guidelines and UnionBank's published card "
        "eligibility requirements. The dataset Income field is treated as annual "
        "PHP income."
    ))
    doc.add_paragraph()
    add_table(doc,
        ["Tier", "Annual Income (PHP)", "Monthly Equivalent", "Typical Products"],
        [
            ["Entry",    "< 360,000",          "< 30,000",   "EON, Blue Card, basic savings"],
            ["Standard", "360,000 – 1,200,000","30k – 100k", "Silver Card, personal loan, auto loan"],
            ["Premium",  "1.2M – 3,000,000",  "100k – 250k","Gold Card, UITF, home loan, Priority Banking"],
            ["Wealth",   "> 3,000,000",         "> 250k",     "Black Card, Wealth Management, GlobalLinker"],
        ],
        col_widths=[1.1, 1.8, 1.5, 2.8]
    )

    doc.add_paragraph()

    # ── 6. Card Upgrade Path ───────────────────────────────────────────────────
    set_heading(doc, "6. Card Upgrade Eligibility", 1, (31, 78, 121))
    add_paragraph(doc, (
        "The recommender follows UnionBank's published card hierarchy: "
        "Blue → Silver → Gold → Black. Upgrade eligibility is checked by comparing "
        "the client's annual income against the minimum income requirement for the "
        "next tier. Only one upgrade step is recommended at a time."
    ))
    doc.add_paragraph()
    add_table(doc,
        ["Upgrade To", "Minimum Annual Income (PHP)", "Monthly Equivalent"],
        [
            ["Silver", "480,000",   "PHP 40,000"],
            ["Gold",   "1,200,000", "PHP 100,000"],
            ["Black",  "3,000,000", "PHP 250,000"],
        ],
        col_widths=[1.5, 2.5, 2.0]
    )

    doc.add_paragraph()

    # ── 7. Rule Layers ─────────────────────────────────────────────────────────
    set_heading(doc, "7. Rule Layers", 1, (31, 78, 121))
    add_paragraph(doc, (
        "Five independent rule layers are evaluated for every client. Results from "
        "all layers are pooled, deduplicated, sorted by priority, and capped at "
        "five recommendations."
    ))

    # 7.1 Retention
    set_heading(doc, "7.1  Retention Rules  (_retention_rules)", 2)
    add_paragraph(doc, (
        "Triggered by elevated attrition probability. Given the model's near-random "
        "performance, these are intentionally conservative — only activating at "
        "probability ≥ 0.45 to avoid flooding every client with retention offers."
    ))
    add_table(doc,
        ["Condition", "Product Recommended", "Priority"],
        [
            ["probability ≥ 0.65",                              "UnionBank Rewards Points Multiplier",  "Critical"],
            ["probability ≥ 0.45 AND tenure ≥ 36 months",       "UnionBank Loyalty Cashback Program",   "High"],
            ["probability ≥ 0.45 AND income_tier = Premium/Wealth", "UnionBank Priority Banking",       "Critical"],
        ],
        col_widths=[2.8, 2.6, 0.8]
    )
    add_paragraph(doc,
        "Note: Retention rules are the weakest layer by design. "
        "The attrition probability is near-random (PR-AUC ≈ 0.05). "
        "These recommendations function as conversation starters for "
        "relationship managers, not deterministic churn signals.",
        italic=True
    )

    doc.add_paragraph()

    # 7.2 Card Upgrade
    set_heading(doc, "7.2  Card Upgrade Rules  (_card_upgrade_rules)", 2)
    add_paragraph(doc, (
        "Objective, income-gated eligibility check. This is the most reliable rule "
        "layer because it is based on verifiable financial data with no dependency "
        "on the weak attrition signal."
    ))
    add_table(doc,
        ["Condition", "Product Recommended", "Priority"],
        [
            ["Current card = Blue AND income ≥ 480k",   "UnionBank Silver Visa Card", "Medium"],
            ["Current card = Silver AND income ≥ 1.2M", "UnionBank Gold Visa Card",   "Medium"],
            ["Current card = Gold AND income ≥ 3M",     "UnionBank Black Visa Card",  "Medium"],
        ],
        col_widths=[2.8, 2.4, 0.8]
    )

    doc.add_paragraph()

    # 7.3 Engagement
    set_heading(doc, "7.3  Engagement Rules  (_engagement_rules)", 2)
    add_paragraph(doc, (
        "Behavioral engagement signals from transaction frequency and credit "
        "utilization. These are the most actionable rules because they are "
        "directly observable and map to specific product interventions."
    ))
    add_table(doc,
        ["Condition", "Signal Meaning", "Product Recommended", "Priority"],
        [
            ["Credit_Utilization < 15%",   "Card underused",          "UnionBank PayDay Promo / Merchant Cashback", "Medium"],
            ["TotalTransactions < 24",     "Digitally disengaged",    "UnionBank Online / EON Digital Banking",     "Medium"],
            ["Credit_Utilization > 85%",   "Near credit limit",       "UnionBank Credit Limit Increase",            "High"],
        ],
        col_widths=[1.8, 1.5, 2.3, 0.8]
    )

    doc.add_paragraph()

    # 7.4 Cross-sell
    set_heading(doc, "7.4  Cross-Sell Rules  (_crosssell_rules)", 2)
    add_paragraph(doc, (
        "Proactive product matching based on income tier, age, and tenure. "
        "Designed to deepen the client relationship and increase switching cost "
        "— both reduce attrition risk without relying on the weak ML signal."
    ))
    add_table(doc,
        ["Condition", "Product Recommended", "Priority"],
        [
            ["income_tier = Standard/Premium/Wealth AND tenure ≥ 12mo",      "UnionBank Personal Loan",       "Low"],
            ["Age 28–50 AND income_tier = Premium/Wealth",                    "UnionBank Home Loan",           "Medium"],
            ["Age 25–45 AND income_tier = Standard/Premium/Wealth",           "UnionBank Auto Loan",           "Low"],
            ["income_tier = Wealth OR (Premium AND utilization < 40%)",        "UnionBank UITF",                "Medium"],
            ["Age ≥ 35 AND income > 1.5M AND income_tier = Premium/Wealth",   "UnionBank Business / GlobalLinker", "Low"],
        ],
        col_widths=[3.0, 2.2, 0.8]
    )

    doc.add_paragraph()

    # 7.5 Life Stage
    set_heading(doc, "7.5  Life Stage Rules  (_lifestage_rules)", 2)
    add_paragraph(doc, (
        "Products aligned to the client's life stage. These are triggered "
        "independently of attrition probability."
    ))
    add_table(doc,
        ["Condition", "Life Stage", "Product Recommended", "Priority"],
        [
            ["Age < 30 AND income_tier = Entry",                          "Young Professional", "UnionBank EON Cyber Account",          "Low"],
            ["Age ≥ 55 AND income_tier = Standard/Premium/Wealth",        "Pre-Retirement",     "UnionBank Time Deposit / BancAssurance","Medium"],
        ],
        col_widths=[2.2, 1.2, 2.2, 0.8]
    )

    doc.add_paragraph()

    # ── 8. Output Structure ────────────────────────────────────────────────────
    set_heading(doc, "8. Output Structure", 1, (31, 78, 121))
    add_paragraph(doc, (
        "Each recommendation contains five fields designed for direct use by a "
        "relationship manager or downstream CRM system."
    ))
    doc.add_paragraph()
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ["product",  "str", "Full UnionBank product name"],
            ["category", "str", "Retention | Upgrade | Engagement | CrossSell | LifeStage"],
            ["reason",   "str", "Plain-language explanation for the RM — references the specific signal that triggered the rule"],
            ["priority", "str", "Critical | High | Medium | Low — determines sort order"],
            ["action",   "str", "Specific next step for the relationship manager or system to execute"],
        ],
        col_widths=[1.1, 1.1, 4.0]
    )
    add_paragraph(doc, (
        "Recommendations are deduplicated (same product cannot appear twice), "
        "sorted by priority (Critical first), and capped at 5 per client."
    ))

    doc.add_paragraph()

    # ── 9. Priority & Deduplication ────────────────────────────────────────────
    set_heading(doc, "9. Priority Ordering & Deduplication", 1, (31, 78, 121))
    add_paragraph(doc, (
        "After all rule layers are evaluated, results are merged into a single list. "
        "The following process is applied before returning:"
    ))
    add_bullet(doc, "Sort by priority: Critical (0) → High (1) → Medium (2) → Low (3)")
    add_bullet(doc, "Deduplicate: if the same product name appears from multiple rules, only the first (highest priority) instance is kept")
    add_bullet(doc, "Cap at 5: only the top 5 unique recommendations are returned per client")
    add_paragraph(doc, (
        "This ensures relationship managers receive a focused, prioritised shortlist "
        "rather than an overwhelming list of every eligible product."
    ))

    doc.add_paragraph()

    # ── 10. Sample Output ──────────────────────────────────────────────────────
    set_heading(doc, "10. Sample API Response (Single Client)", 1, (31, 78, 121))
    add_paragraph(doc, "Input: 47-year-old male, PHP 1.6M income, Gold card, 3 transactions, 46 months tenure.")
    add_code_block(doc,
        '{\n'
        '  "attrition_probability": 0.0521,\n'
        '  "attrition_flag": 0,\n'
        '  "attrition_risk_tier": "Low",\n'
        '  "client_segment": "Mid Career · Premium",\n'
        '  "top_drivers": {\n'
        '    "Avg_Txn_per_Tenure_Year": -0.1832,\n'
        '    "TotalTransactions": -0.1204,\n'
        '    "Age": 0.0871,\n'
        '    "Income": -0.0634,\n'
        '    "Credit_Utilization": -0.0512\n'
        '  },\n'
        '  "recommended_products": [\n'
        '    {\n'
        '      "product": "UnionBank Black Visa Card",\n'
        '      "category": "Upgrade",\n'
        '      "reason": "Income qualifies for Black tier (PHP 1,600,000/yr)",\n'
        '      "priority": "Medium",\n'
        '      "action": "Initiate card upgrade offer to UnionBank Black Visa."\n'
        '    },\n'
        '    {\n'
        '      "product": "UnionBank Online / EON Digital Banking",\n'
        '      "category": "Engagement",\n'
        '      "reason": "Only 3 total transactions recorded — less than 1 per month",\n'
        '      "priority": "Medium",\n'
        '      "action": "Send digital onboarding push notification."\n'
        '    },\n'
        '    {\n'
        '      "product": "UnionBank Unit Investment Trust Fund (UITF)",\n'
        '      "category": "CrossSell",\n'
        '      "reason": "High-income client with moderate card utilization",\n'
        '      "priority": "Medium",\n'
        '      "action": "Refer to Trust & Investment Division."\n'
        '    }\n'
        '  ]\n'
        '}'
    )

    doc.add_paragraph()

    # ── 11. Limitations ────────────────────────────────────────────────────────
    set_heading(doc, "11. Known Limitations & Future Improvements", 1, (31, 78, 121))

    set_heading(doc, "11.1  Current Limitations", 2)
    add_bullet(doc, "Income thresholds are fixed constants — not dynamically calibrated to portfolio-level data")
    add_bullet(doc, "Attrition probability is near-random (PR-AUC ≈ 0.05); retention recommendations are soft signals only")
    add_bullet(doc, "No personalization beyond the 11 available features — product affinity scoring not possible without transaction history")
    add_bullet(doc, "Card eligibility requirements are approximate — actual bank policy may include additional criteria (credit score, existing delinquency, etc.)")

    set_heading(doc, "11.2  Recommended Enhancements", 2)
    add_bullet(doc, "Integrate transaction-level behavioral features (recency, frequency, trend slope) to improve both attrition signal and recommendation relevance")
    add_bullet(doc, "Replace static income thresholds with dynamic percentile-based segmentation recalculated at each model refresh cycle")
    add_bullet(doc, "Add product acceptance feedback loop — log which recommendations led to conversions and use this to weight rules over time")
    add_bullet(doc, "Add a product exclusion list — do not recommend products the client already holds")

    doc.add_paragraph()

    # ── 12. File Reference ─────────────────────────────────────────────────────
    set_heading(doc, "12. Code Reference", 1, (31, 78, 121))
    add_table(doc,
        ["File", "Purpose"],
        [
            ["api/recommender.py",      "All rule logic — get_recommendations(), five rule layer functions, RecommendationBundle"],
            ["api/schemas.py",          "Pydantic models — RecommendedProduct, PredictionResponse (includes recommended_products field)"],
            ["api/transformers.py",     "Custom sklearn transformers — WinsorizationTransformer, LRFeatureEngineeringTransformer"],
            ["main.py",                 "FastAPI app — calls get_recommendations() after pipeline.predict_proba()"],
            ["scripts/serialize_model.py", "Training script — builds and saves the inference pipeline to models/"],
        ],
        col_widths=[2.2, 4.0]
    )

    # ── Save ───────────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"Document saved: {OUT_PATH}")


if __name__ == "__main__":
    build_document()
